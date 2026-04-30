# report.py
# Генерация Word-отчёта по шаблону template.docx.
# Этот файл открывает шаблон, заменяет метки на расчёты, таблицы, анализ и вставляет графики.

from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from variants import AMPLITUDE_DISTRIBUTION_NAMES


# -----------------------------
# Форматирование чисел
# -----------------------------

def ru_number(value, digits=2):
    """Число с запятой вместо точки."""
    try:
        text = f"{value:.{digits}f}"
    except Exception:
        text = str(value)
    return text.replace(".", ",")


def ru_sci(value, digits=3):
    """Научная запись с запятой вместо точки."""
    text = f"{value:.{digits}e}"
    return text.replace(".", ",")


def format_F(value):
    """Красивое отображение вероятности ложной тревоги."""
    if value == 1e-4:
        return "10^-4"
    if value == 1e-5:
        return "10^-5"
    if value == 1e-6:
        return "10^-6"
    if value == 1e-7:
        return "10^-7"
    if value == 1e-8:
        return "10^-8"
    if value == 1e-9:
        return "10^-9"
    return str(value).replace(".", ",")


# -----------------------------
# Низкоуровневые функции вставки в Word
# -----------------------------

def find_paragraph_with_text(document, marker):
    """Поиск абзаца, содержащего метку."""
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            return paragraph
    return None


def replace_marker_with_text(document, marker, text):
    """Замена метки на текст."""
    paragraph = find_paragraph_with_text(document, marker)
    if paragraph is None:
        return

    paragraph.text = paragraph.text.replace(marker, text)


def insert_paragraph_after(paragraph, text="", style=None):
    """Вставка нового абзаца после существующего."""
    new_paragraph = paragraph.insert_paragraph_before(text, style=style)
    paragraph._p.addprevious(new_paragraph._p)
    return new_paragraph


def add_picture_after_paragraph(document, paragraph, image_path, width_inches=6.2):
    """Вставка изображения сразу после указанного абзаца."""
    image_path = str(image_path)

    new_paragraph = document.add_paragraph()
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_paragraph.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def add_table_after_paragraph(document, paragraph, rows):
    """
    Вставка таблицы после абзаца.

    rows — список списков:
    [
        ["Параметр", "Обозначение", "Значение"],
        ...
    ]
    """
    table = document.add_table(rows=len(rows), cols=len(rows[0]))

    # В разных версиях Word стиль таблицы может называться по-разному.
    # Поэтому пробуем несколько вариантов, а если не получилось — оставляем стиль по умолчанию.
    for style_name in ("Table Grid", "Сетка таблицы"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)

    paragraph._p.addnext(table._tbl)
    return table


def replace_marker_with_table(document, marker, rows):
    """Замена метки на таблицу."""
    paragraph = find_paragraph_with_text(document, marker)
    if paragraph is None:
        return

    paragraph.text = paragraph.text.replace(marker, "")
    add_table_after_paragraph(document, paragraph, rows)


def replace_marker_with_picture(document, marker, image_path, caption):
    """Замена метки на подпись и изображение."""
    paragraph = find_paragraph_with_text(document, marker)
    if paragraph is None:
        return

    paragraph.text = caption
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture_after_paragraph(document, paragraph, image_path)


# -----------------------------
# Таблица исходных данных
# -----------------------------

def build_input_data_rows(variant_number, data):
    """Формирует строки таблицы исходных данных."""
    distribution_name = AMPLITUDE_DISTRIBUTION_NAMES[data["amplitude_distribution"]]

    delta_p = data.get("delta_p")
    delta_p_text = "-" if delta_p is None else str(delta_p).replace(".", ",")

    return [
        ["Параметр", "Обозначение", "Значение"],
        ["Номер варианта", "—", variant_number],
        ["Рабочая длина волны", "λв", f"{str(data['lambda_v']).replace('.', ',')} м"],
        ["Размер АФАР по горизонтали", "Lax", f"{str(data['Lax']).replace('.', ',')} м"],
        ["Размер АФАР по вертикали", "Lay", f"{str(data['Lay']).replace('.', ',')} м"],
        ["Угол наклона полотна АФАР", "γa", f"{data['gamma_a']}°"],
        ["Высота подъёма АФАР", "ha", f"{str(data['ha']).replace('.', ',')} м"],
        ["Амплитудное распределение", "АР", distribution_name],
        ["Пьедестал", "Δп", delta_p_text],
        ["КПД антенно-фидерных трактов", "η", str(data["eta"]).replace(".", ",")],
        ["Импульсная мощность одного модуля", "Pи1", f"{data['Pi1']} Вт"],
        ["Вероятность правильного обнаружения", "D", str(data["D"]).replace(".", ",")],
        ["Вероятность ложной тревоги", "F", format_F(data["F"])],
        ["Длительность импульса", "τи", f"{str(data['tau_i_us']).replace('.', ',')} мкс"],
        ["Ширина спектра импульса", "Δfи", f"{str(data['delta_f_i_mhz']).replace('.', ',')} МГц"],
        ["Коэффициент потерь", "kпот", f"{str(data['k_loss_db']).replace('.', ',')} дБ"],
        ["ЭПР цели", "σц", f"{str(data['sigma_c']).replace('.', ',')} м²"],
        ["Высота цели", "hц", f"{str(data['h_c']).replace('.', ',')} м"],
        ["Мощность активной помехи", "Pп", f"{str(data['P_p']).replace('.', ',')} Вт"],
        ["Ширина спектра активной помехи", "Δfп", f"{str(data['delta_f_p_mhz']).replace('.', ',')} МГц"],
        ["Коэффициент учёта поляризации", "kпол", str(data["k_pol"]).replace(".", ",")],
        ["Коэффициент усиления антенны ПАП", "Gап", f"{str(data['G_ap_db']).replace('.', ',')} дБ"],
        ["Расстояние от РЛС до ПАП", "rп", f"{str(data['r_p_km']).replace('.', ',')} км"],
        ["Азимут ПАП", "αп", f"{data['alpha_p']}°"],
        ["Угол места ПАП", "εп", f"{data['epsilon_p']}°"],
    ]


# -----------------------------
# Расчётные тексты по пунктам
# -----------------------------

def build_calculation_texts(data, results):
    """Формирует тексты расчётов для меток POINT_1_CALC ... POINT_18_CALC."""
    distribution_name = AMPLITUDE_DISTRIBUTION_NAMES[data["amplitude_distribution"]]

    texts = {}

    texts["POINT_1_CALC"] = f"""По исходным данным рабочая длина волны РЛС составляет:

λв = {str(data['lambda_v']).replace('.', ',')} м.

Рабочая частота определяется выражением:

f0 = c / λв.

Подставляя исходные данные, получим:

f0 = 3·10^8 / {str(data['lambda_v']).replace('.', ',')} = {ru_sci(results['f0'], 2)} Гц = {ru_number(results['f0_ghz'], 2)} ГГц.

Так как λв = {str(data['lambda_v']).replace('.', ',')} м, РЛС работает в {results['wave_range']} диапазоне волн."""

    texts["POINT_2_CALC"] = f"""Шаг прямоугольной решётки принимается равным половине длины волны:

dx = dy = 0,5·λв = 0,5·{str(data['lambda_v']).replace('.', ',')} = {ru_number(results['dx'], 3)} м.

Количество излучателей по горизонтали:

Nx = Lax / dx = {str(data['Lax']).replace('.', ',')} / {ru_number(results['dx'], 3)} = {ru_number(results['Nx_float'], 2)}.

С учётом округления в меньшую сторону:

Nx = {results['Nx']}.

Количество излучателей по вертикали:

Ny = Lay / dy = {str(data['Lay']).replace('.', ',')} / {ru_number(results['dy'], 3)} = {ru_number(results['Ny_float'], 2)}.

С учётом округления в меньшую сторону:

Ny = {results['Ny']}.

Общее количество излучателей:

NΣ = Nx·Ny = {results['Nx']}·{results['Ny']} = {results['N_sum']}."""

    texts["POINT_3_CALC"] = f"""Геометрическая площадь АФАР:

Sаг = Lax·Lay = {str(data['Lax']).replace('.', ',')}·{str(data['Lay']).replace('.', ',')} = {ru_number(results['S_geo'], 3)} м².

Для амплитудного распределения «{distribution_name}» коэффициент использования поверхности:

kип = kипx·kипy = {ru_number(results['k_ip_x'], 3)}·{ru_number(results['k_ip_y'], 3)} = {ru_number(results['k_ip'], 3)}.

Эффективная площадь АФАР:

Sa = Sаг·kип = {ru_number(results['S_geo'], 3)}·{ru_number(results['k_ip'], 3)} = {ru_number(results['S_a'], 3)} м².

Коэффициент усиления АФАР:

Ga = 4π·Sa·η / λв² = 4π·{ru_number(results['S_a'], 3)}·{str(data['eta']).replace('.', ',')} / {str(data['lambda_v']).replace('.', ',')}² = {ru_number(results['G_a'], 2)}.

Ga[дБ] = 10lg(Ga) = 10lg({ru_number(results['G_a'], 2)}) = {ru_number(results['G_a_db'], 2)} дБ.

Ширина главного луча в азимутальной плоскости:

ΔΘα = {ru_number(results['beamwidth_constant'], 1)}·λв / Lax = {ru_number(results['beamwidth_constant'], 1)}·{str(data['lambda_v']).replace('.', ',')} / {str(data['Lax']).replace('.', ',')} = {ru_number(results['delta_theta_alpha'], 2)}°.

Ширина главного луча в угломестной плоскости:

ΔΘε = {ru_number(results['beamwidth_constant'], 1)}·λв / Lay = {ru_number(results['beamwidth_constant'], 1)}·{str(data['lambda_v']).replace('.', ',')} / {str(data['Lay']).replace('.', ',')} = {ru_number(results['delta_theta_epsilon'], 2)}°.

Уровень боковых лепестков:

ηбл = {ru_number(results['eta_bl_db'], 1)} дБ."""

    texts["POINT_4_CALC"] = f"""Для анализа направленных свойств АФАР построены нормированные диаграммы направленности в азимутальной и угломестной плоскостях.

Нормированная диаграмма направленности определяется выражением:

FN(θ,φ) = F(θ,φ) / F(θ0,φ0).

В логарифмическом масштабе:

FN(θ,φ)[дБ] = 20lg(FN(θ,φ)).

Для заданных исходных данных построены графики FN(θ,0) и FN(θ,90)."""

    texts["POINT_5_CALC"] = f"""Ширина главного луча определяется по уровню -3 дБ.

По расчётам:

ΔΘα = {ru_number(results['delta_theta_alpha'], 2)}°,

ΔΘε = {ru_number(results['delta_theta_epsilon'], 2)}°.

На построенных диаграммах направленности ширина главного луча определяется по пересечению графика с уровнем -3 дБ.

Для амплитудного распределения «{distribution_name}» теоретический уровень боковых лепестков составляет:

ηбл = {ru_number(results['eta_bl_db'], 1)} дБ.

На графиках дополнительно отмечена горизонтальная линия уровня боковых лепестков."""

    texts["POINT_6_CALC"] = f"""Угол наклона полотна АФАР относительно земной поверхности:

γa = {data['gamma_a']}°.

Угол наклона нормали к АФАР определяется выражением:

γε норм = 90° - γa.

Подставляя значение:

γε норм = 90° - {data['gamma_a']}° = {ru_number(results['gamma_epsilon_norm'], 2)}°."""

    texts["POINT_7_CALC"] = f"""Импульсная мощность РЛС определяется как произведение мощности одного модуля на количество модулей:

Pи = Pи1·NΣ.

Подставляя значения:

Pи = {data['Pi1']}·{results['N_sum']} = {ru_number(results['P_i'], 0)} Вт = {ru_number(results['P_i_kw'], 2)} кВт."""

    texts["POINT_8_CALC"] = f"""Средняя мощность РЛС определяется через скважность:

Pср = Pи / Q.

При Q = 8:

Pср = {ru_number(results['P_i'], 0)} / 8 = {ru_number(results['P_avg'], 0)} Вт = {ru_number(results['P_avg_kw'], 2)} кВт."""

    texts["POINT_9_CALC"] = f"""Коэффициент шума в линейном виде:

kш = 10^(2/10) = {ru_number(results['k_noise'], 3)}.

Эффективная полоса пропускания приёмника:

ΔFпр = 1 / τи = 1 / ({str(data['tau_i_us']).replace('.', ',')}·10^-6) = {ru_number(results['delta_F_pr'], 1)} Гц.

Спектральная плотность мощности шума:

N0 = kБ·kш·T0 = 1,38·10^-23·{ru_number(results['k_noise'], 3)}·290 = {ru_sci(results['N0'], 3)} Вт/Гц.

Чувствительность приёмного устройства:

Pпр min = N0·ΔFпр = {ru_sci(results['N0'], 3)}·{ru_number(results['delta_F_pr'], 1)} = {ru_sci(results['P_pr_min'], 3)} Вт."""

    texts["POINT_10_CALC"] = f"""Пороговое отношение сигнал/шум определяется выражением:

q² = 2·(ln(F) / ln(D) - 1).

Подставляя значения:

q² = 2·(ln({format_F(data['F'])}) / ln({str(data['D']).replace('.', ',')}) - 1) = {ru_number(results['q2'], 2)}.

В децибелах:

q²[дБ] = 10lg(q²) = 10lg({ru_number(results['q2'], 2)}) = {ru_number(results['q2_db'], 2)} дБ.

Пороговая чувствительность:

Pпор = q²·Pпр min = {ru_number(results['q2'], 2)}·{ru_sci(results['P_pr_min'], 3)} = {ru_sci(results['P_threshold'], 3)} Вт."""

    texts["POINT_11_CALC"] = f"""Минимальная дальность действия импульсной РЛС определяется длительностью зондирующего импульса:

Rmin = c·τи / 2.

Подставляя значения:

Rmin = 3·10^8·{str(data['tau_i_us']).replace('.', ',')}·10^-6 / 2 = {ru_number(results['R_min'], 0)} м = {ru_number(results['R_min_km'], 2)} км."""

    texts["POINT_12_CALC"] = f"""Максимальная дальность действия РЛС при неотклонённом главном луче определяется выражением:

Rmax = [Pи·Ga²·λв²·σц / ((4π)³·q²·N0·ΔFпр·kпот)]^(1/4).

Коэффициент потерь в линейном виде:

kпот = 10^({str(data['k_loss_db']).replace('.', ',')}/10) = {ru_number(results['k_loss'], 3)}.

Подставляя значения:

Rmax = [{ru_number(results['P_i'], 0)}·{ru_number(results['G_a'], 2)}²·{str(data['lambda_v']).replace('.', ',')}²·{str(data['sigma_c']).replace('.', ',')} / ((4π)³·{ru_number(results['q2'], 2)}·{ru_sci(results['N0'], 3)}·{ru_number(results['delta_F_pr'], 1)}·{ru_number(results['k_loss'], 3)})]^(1/4).

Rmax = {ru_number(results['R_max'], 0)} м = {ru_number(results['R_max_km'], 2)} км."""

    texts["POINT_13_CALC"] = f"""При сканировании зоны обзора максимальная дальность уменьшается с учётом отклонения главного луча от нормали к АФАР.

Для зависимости от азимута при ε = γε норм:

Rmax(α, γε норм) = Rmax·√(cos α).

При α = 0°:

Rmax(0°, {ru_number(results['gamma_epsilon_norm'], 2)}°) = {ru_number(results['R_max_km'], 2)} км.

При α = ±60°:

Rmax(±60°, {ru_number(results['gamma_epsilon_norm'], 2)}°) = {ru_number(results['R_max_km'], 2)}·√(cos 60°) = {ru_number(results['R_max_alpha_edge_km'], 2)} км.

Для зависимости от угла места при α = 0:

Rmax(0, ε) = Rmax·√(cos(ε - γε норм)).

При ε = 5°:

Rmax(0, 5°) = {ru_number(results['R_max_km'], 2)}·√(cos(5° - {ru_number(results['gamma_epsilon_norm'], 2)}°)) = {ru_number(results['R_max_epsilon_min_km'], 2)} км.

При ε = 70°:

Rmax(0, 70°) = {ru_number(results['R_max_km'], 2)}·√(cos(70° - {ru_number(results['gamma_epsilon_norm'], 2)}°)) = {ru_number(results['R_max_epsilon_max_km'], 2)} км."""

    texts["POINT_14_CALC"] = f"""Разрешающая способность по дальности:

δR = c / (2·Δfи).

Подставляя значения:

δR = 3·10^8 / (2·{str(data['delta_f_i_mhz']).replace('.', ',')}·10^6) = {ru_number(results['delta_R'], 2)} м.

Разрешающая способность по радиальной скорости:

δVr = λв / (2·τи).

δVr = {str(data['lambda_v']).replace('.', ',')} / (2·{str(data['tau_i_us']).replace('.', ',')}·10^-6) = {ru_number(results['delta_Vr'], 2)} м/с.

Разрешающая способность по азимуту:

δΘα = ΔΘα = {ru_number(results['delta_theta_alpha'], 2)}° = {ru_number(results['delta_theta_alpha_min'], 1)}′.

Разрешающая способность по углу места:

δΘε = ΔΘε = {ru_number(results['delta_theta_epsilon'], 2)}° = {ru_number(results['delta_theta_epsilon_min'], 1)}′."""

    texts["POINT_15_CALC"] = f"""Точность измерения дальности:

ξR = δR / (π·q²).

ξR = {ru_number(results['delta_R'], 2)} / (π·{ru_number(results['q2'], 2)}) = {ru_number(results['xi_R'], 3)} м.

Точность измерения радиальной скорости:

ξVr = δVr / (π·q²).

ξVr = {ru_number(results['delta_Vr'], 2)} / (π·{ru_number(results['q2'], 2)}) = {ru_number(results['xi_Vr'], 3)} м/с.

Точность измерения азимута:

ξΘα = δΘα / (π·q²).

ξΘα = {ru_number(results['delta_theta_alpha'], 2)} / (π·{ru_number(results['q2'], 2)}) = {ru_number(results['xi_theta_alpha'], 4)}° = {ru_number(results['xi_theta_alpha_min'], 2)}′.

Точность измерения угла места:

ξΘε = δΘε / (π·q²).

ξΘε = {ru_number(results['delta_theta_epsilon'], 2)} / (π·{ru_number(results['q2'], 2)}) = {ru_number(results['xi_theta_epsilon'], 4)}° = {ru_number(results['xi_theta_epsilon_min'], 2)}′."""

    texts["POINT_16_CALC"] = f"""Для маловысотной цели максимальная дальность обнаружения определяется выражением:

Rmax МВЦ = √((4π·ha·hц / λв)·Rmax).

Подставляя значения:

Rmax МВЦ = √((4π·{str(data['ha']).replace('.', ',')}·{str(data['h_c']).replace('.', ',')} / {str(data['lambda_v']).replace('.', ',')})·{ru_number(results['R_max'], 0)}).

Rmax МВЦ = {ru_number(results['R_max_mvc'], 0)} м = {ru_number(results['R_max_mvc_km'], 2)} км."""

    texts["POINT_17_CALC"] = f"""При воздействии активных помех рассматриваются два случая: самоприкрытие и внешнее прикрытие.

Для направления α = 0°, ε = γε норм = {ru_number(results['gamma_epsilon_norm'], 2)}° получены следующие значения:

Без помех:

Rmax = {ru_number(results['R_max_km'], 2)} км.

Самоприкрытие при воздействии помехи в направлении главного луча:

Rmax п сп = {ru_number(results['R_max_p_sp_km'], 2)} км.

Самоприкрытие при воздействии помехи через боковые лепестки:

Rmax п бл сп = {ru_number(results['R_max_p_bl_sp_km'], 2)} км.

Внешнее прикрытие при воздействии помехи в направлении главного луча:

Rmax п вп = {ru_number(results['R_max_p_vp_km'], 2)} км.

Внешнее прикрытие при воздействии помехи через боковые лепестки:

Rmax п бл вп = {ru_number(results['R_max_p_bl_vp_km'], 2)} км."""

    texts["POINT_18_CALC"] = f"""Коэффициент сжатия определяется как отношение максимальной дальности при воздействии помехи к максимальной дальности без помех:

kсж = Rmax п / Rmax.

Для самоприкрытия в направлении главного луча:

kсж п сп = Rmax п сп / Rmax = {ru_number(results['R_max_p_sp_km'], 2)} / {ru_number(results['R_max_km'], 2)} = {ru_number(results['k_szh_p_sp'], 3)}.

Для самоприкрытия через боковые лепестки:

kсж п бл сп = Rmax п бл сп / Rmax = {ru_number(results['R_max_p_bl_sp_km'], 2)} / {ru_number(results['R_max_km'], 2)} = {ru_number(results['k_szh_p_bl_sp'], 3)}.

Для внешнего прикрытия в направлении главного луча:

kсж п вп = Rmax п вп / Rmax = {ru_number(results['R_max_p_vp_km'], 2)} / {ru_number(results['R_max_km'], 2)} = {ru_number(results['k_szh_p_vp'], 3)}.

Для внешнего прикрытия через боковые лепестки:

kсж п бл вп = Rmax п бл вп / Rmax = {ru_number(results['R_max_p_bl_vp_km'], 2)} / {ru_number(results['R_max_km'], 2)} = {ru_number(results['k_szh_p_bl_vp'], 3)}."""

    texts["CONCLUSION_CALC"] = f"""В ходе выполнения курсовой работы была проведена оценка информационных возможностей секторной импульсной РЛС с плоской прямоугольной АФАР.

Рабочая длина волны РЛС составляет {str(data['lambda_v']).replace('.', ',')} м, что соответствует {results['wave_range']} диапазону волн. Общее количество излучателей АФАР составило {results['N_sum']}. Максимальная дальность действия РЛС по цели с ЭПР {str(data['sigma_c']).replace('.', ',')} м² при неотклонённом главном луче составила {ru_number(results['R_max_km'], 2)} км.

Для маловысотной цели дальность обнаружения уменьшилась до {ru_number(results['R_max_mvc_km'], 2)} км. При воздействии активных помех наибольшее снижение дальности наблюдается при самоприкрытии в направлении главного луча, где дальность уменьшается до {ru_number(results['R_max_p_sp_km'], 2)} км, а коэффициент сжатия зоны действия составляет {ru_number(results['k_szh_p_sp'], 3)}.

Таким образом, рассматриваемая РЛС обладает достаточно высокой дальностью обнаружения при отсутствии помех, однако её возможности существенно снижаются при обнаружении маловысотных целей и при воздействии активных помех в направлении главного луча."""

    return texts


# -----------------------------
# Итоговая таблица и анализ результатов
# -----------------------------

def build_results_table_rows(results):
    """Формирует строки итоговой таблицы результатов."""
    return [
        ["Раздел", "Параметр", "Обозначение", "Значение", "Единица"],

        ["Рабочий диапазон", "Рабочая частота", "f0", ru_number(results["f0_ghz"], 2), "ГГц"],
        ["Рабочий диапазон", "Диапазон волн", "—", results["wave_range"], "—"],

        ["АФАР", "Шаг решётки по горизонтали", "dx", ru_number(results["dx"], 4), "м"],
        ["АФАР", "Шаг решётки по вертикали", "dy", ru_number(results["dy"], 4), "м"],
        ["АФАР", "Количество излучателей по горизонтали", "Nx", results["Nx"], "—"],
        ["АФАР", "Количество излучателей по вертикали", "Ny", results["Ny"], "—"],
        ["АФАР", "Общее количество излучателей", "NΣ", results["N_sum"], "—"],
        ["АФАР", "Геометрическая площадь АФАР", "Sаг", ru_number(results["S_geo"], 3), "м²"],
        ["АФАР", "Эффективная площадь АФАР", "Sa", ru_number(results["S_a"], 3), "м²"],
        ["АФАР", "Коэффициент усиления", "Ga", ru_number(results["G_a"], 2), "—"],
        ["АФАР", "Коэффициент усиления", "Ga", ru_number(results["G_a_db"], 2), "дБ"],
        ["АФАР", "Ширина луча по азимуту", "ΔΘα", ru_number(results["delta_theta_alpha"], 2), "град"],
        ["АФАР", "Ширина луча по углу места", "ΔΘε", ru_number(results["delta_theta_epsilon"], 2), "град"],
        ["АФАР", "Уровень боковых лепестков", "ηбл", ru_number(results["eta_bl_db"], 1), "дБ"],
        ["АФАР", "Угол нормали к АФАР", "γε норм", ru_number(results["gamma_epsilon_norm"], 2), "град"],

        ["Мощность", "Импульсная мощность РЛС", "Pи", ru_number(results["P_i_kw"], 2), "кВт"],
        ["Мощность", "Средняя мощность РЛС", "Pср", ru_number(results["P_avg_kw"], 2), "кВт"],

        ["Приёмник", "Эффективная полоса пропускания", "ΔFпр", ru_number(results["delta_F_pr"], 1), "Гц"],
        ["Приёмник", "Чувствительность приёмника", "Pпр min", ru_sci(results["P_pr_min"], 3), "Вт"],
        ["Приёмник", "Пороговое отношение сигнал/шум", "q²", ru_number(results["q2"], 2), "—"],
        ["Приёмник", "Пороговое отношение сигнал/шум", "q²", ru_number(results["q2_db"], 2), "дБ"],
        ["Приёмник", "Пороговая чувствительность", "Pпор", ru_sci(results["P_threshold"], 3), "Вт"],

        ["Дальность", "Минимальная дальность", "Rmin", ru_number(results["R_min_km"], 2), "км"],
        ["Дальность", "Максимальная дальность", "Rmax", ru_number(results["R_max_km"], 2), "км"],
        ["Дальность", "Rmax при α = ±60°", "Rmax(±60°)", ru_number(results["R_max_alpha_edge_km"], 2), "км"],
        ["Дальность", "Rmax при ε = 5°", "Rmax(5°)", ru_number(results["R_max_epsilon_min_km"], 2), "км"],
        ["Дальность", "Rmax при ε = 70°", "Rmax(70°)", ru_number(results["R_max_epsilon_max_km"], 2), "км"],
        ["Дальность", "Дальность обнаружения маловысотной цели", "Rmax МВЦ", ru_number(results["R_max_mvc_km"], 2), "км"],

        ["Разрешающая способность", "По дальности", "δR", ru_number(results["delta_R"], 2), "м"],
        ["Разрешающая способность", "По радиальной скорости", "δVr", ru_number(results["delta_Vr"], 2), "м/с"],
        ["Разрешающая способность", "По азимуту", "δΘα", ru_number(results["delta_theta_alpha"], 2), "град"],
        ["Разрешающая способность", "По углу места", "δΘε", ru_number(results["delta_theta_epsilon"], 2), "град"],

        ["Точность измерения", "По дальности", "ξR", ru_number(results["xi_R"], 3), "м"],
        ["Точность измерения", "По радиальной скорости", "ξVr", ru_number(results["xi_Vr"], 3), "м/с"],
        ["Точность измерения", "По азимуту", "ξΘα", ru_number(results["xi_theta_alpha"], 4), "град"],
        ["Точность измерения", "По углу места", "ξΘε", ru_number(results["xi_theta_epsilon"], 4), "град"],

        ["Помехи", "Самоприкрытие, главный луч", "Rmax п сп", ru_number(results["R_max_p_sp_km"], 2), "км"],
        ["Помехи", "Самоприкрытие, боковые лепестки", "Rmax п бл сп", ru_number(results["R_max_p_bl_sp_km"], 2), "км"],
        ["Помехи", "Внешнее прикрытие, главный луч", "Rmax п вп", ru_number(results["R_max_p_vp_km"], 2), "км"],
        ["Помехи", "Внешнее прикрытие, боковые лепестки", "Rmax п бл вп", ru_number(results["R_max_p_bl_vp_km"], 2), "км"],

        ["Коэффициенты сжатия", "Самоприкрытие, главный луч", "kсж п сп", ru_number(results["k_szh_p_sp"], 3), "—"],
        ["Коэффициенты сжатия", "Самоприкрытие, боковые лепестки", "kсж п бл сп", ru_number(results["k_szh_p_bl_sp"], 3), "—"],
        ["Коэффициенты сжатия", "Внешнее прикрытие, главный луч", "kсж п вп", ru_number(results["k_szh_p_vp"], 3), "—"],
        ["Коэффициенты сжатия", "Внешнее прикрытие, боковые лепестки", "kсж п бл вп", ru_number(results["k_szh_p_bl_vp"], 3), "—"],
    ]


def build_results_analysis_text(results):
    """Формирует автоматический анализ результатов для Word-отчёта."""
    rmax = results["R_max_km"]
    rmin = results["R_min_km"]
    r_mvc = results["R_max_mvc_km"]

    interference_cases = [
        ("самоприкрытие в направлении главного луча", results["R_max_p_sp_km"], results["k_szh_p_sp"]),
        ("самоприкрытие через боковые лепестки", results["R_max_p_bl_sp_km"], results["k_szh_p_bl_sp"]),
        ("внешнее прикрытие в направлении главного луча", results["R_max_p_vp_km"], results["k_szh_p_vp"]),
        ("внешнее прикрытие через боковые лепестки", results["R_max_p_bl_vp_km"], results["k_szh_p_bl_vp"]),
    ]

    worst_name, worst_range, worst_k = min(interference_cases, key=lambda item: item[2])

    if worst_k > 0:
        compression_times = 1 / worst_k
    else:
        compression_times = 0

    mvc_ratio = r_mvc / rmax if rmax > 0 else 0
    mvc_drop_percent = (1 - mvc_ratio) * 100

    return f"""По заданным исходным данным РЛС работает в {results['wave_range']} диапазоне волн. Рабочая частота составляет {ru_number(results['f0_ghz'], 2)} ГГц.

В составе АФАР используется {results['N_sum']} излучателей. Расчётная ширина главного луча составляет {ru_number(results['delta_theta_alpha'], 2)}° в азимутальной плоскости и {ru_number(results['delta_theta_epsilon'], 2)}° в угломестной плоскости.

Минимальная дальность действия РЛС равна {ru_number(rmin, 2)} км, а максимальная дальность обнаружения цели при неотклонённом главном луче составляет {ru_number(rmax, 2)} км.

Для маловысотной цели максимальная дальность уменьшается до {ru_number(r_mvc, 2)} км. Это составляет примерно {ru_number(mvc_ratio, 2)} от максимальной дальности без учёта малой высоты цели, то есть снижение составляет около {ru_number(mvc_drop_percent, 1)}%.

Наиболее опасным режимом воздействия активной помехи является {worst_name}. В этом случае дальность обнаружения уменьшается до {ru_number(worst_range, 2)} км, а коэффициент сжатия зоны действия составляет {ru_number(worst_k, 3)}. Это означает, что зона действия по дальности уменьшается примерно в {ru_number(compression_times, 1)} раза.

Таким образом, при отсутствии помех РЛС обеспечивает значительную дальность обнаружения, однако при воздействии активной помехи в наиболее неблагоприятном направлении зона действия может существенно сокращаться."""


# -----------------------------
# Графики
# -----------------------------

GRAPH_CAPTIONS = {
    "GRAPH_DN_AZIMUTH": "Рисунок 1 — Нормированная диаграмма направленности АФАР в азимутальной плоскости.",
    "GRAPH_DN_ELEVATION": "Рисунок 2 — Нормированная диаграмма направленности АФАР в угломестной плоскости.",
    "GRAPH_RMAX_ALPHA": "Рисунок 3 — Зависимость максимальной дальности действия РЛС от азимута.",
    "GRAPH_RMAX_EPSILON": "Рисунок 4 — Зависимость максимальной дальности действия РЛС от угла места.",
    "GRAPH_SELF_ALPHA": "Рисунок 5 — Зависимости максимальной дальности действия РЛС от азимута при самоприкрытии.",
    "GRAPH_SELF_EPSILON": "Рисунок 6 — Зависимости максимальной дальности действия РЛС от угла места при самоприкрытии.",
    "GRAPH_EXTERNAL_ALPHA": "Рисунок 7 — Зависимости максимальной дальности действия РЛС от азимута при внешнем прикрытии.",
    "GRAPH_EXTERNAL_EPSILON": "Рисунок 8 — Зависимости максимальной дальности действия РЛС от угла места при внешнем прикрытии.",
}


# -----------------------------
# Основная функция создания отчёта
# -----------------------------

def create_report(
    template_path,
    output_path,
    variant_number,
    student_name,
    group_name,
    data,
    results,
    graph_paths,
):
    """
    Создаёт Word-отчёт по шаблону.

    template_path — путь к template.docx;
    output_path — куда сохранить готовый отчёт;
    graph_paths — словарь с путями к графикам из create_all_plots().
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(exist_ok=True)

    document = Document(str(template_path))

    replace_marker_with_text(document, "{{VARIANT_NUMBER}}", str(variant_number))
    replace_marker_with_text(document, "{{STUDENT_NAME}}", student_name)
    replace_marker_with_text(document, "{{GROUP_NAME}}", group_name)

    input_rows = build_input_data_rows(variant_number, data)
    replace_marker_with_table(document, "{{INPUT_DATA_TABLE}}", input_rows)

    results_rows = build_results_table_rows(results)
    replace_marker_with_table(document, "{{RESULTS_TABLE}}", results_rows)

    calculation_texts = build_calculation_texts(data, results)

    for key, text in calculation_texts.items():
        marker = "{{" + key + "}}"
        replace_marker_with_text(document, marker, text)

    for graph_key, image_path in graph_paths.items():
        marker = "{{" + graph_key + "}}"
        caption = GRAPH_CAPTIONS.get(graph_key, graph_key)
        replace_marker_with_picture(document, marker, image_path, caption)

    document.save(str(output_path))

    return output_path
